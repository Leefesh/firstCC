"""Convert 人工智能概论期末项目说明.md to a formatted Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style helpers ──
def set_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Set East-Asian font
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def set_line_spacing(paragraph, multiple=1.5):
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.makeelement(qn('w:spacing'), {})
    spacing.set(qn('w:line'), str(int(multiple * 240)))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_heading_line(text, level=1):
    """Add a heading with proper formatting."""
    p = doc.add_paragraph()
    set_line_spacing(p, 1.5)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_font(run, size=16, bold=True)  # 三号 ≈ 16pt
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_font(run, size=14, bold=True)  # 四号 ≈ 14pt
    return p

def add_body(text):
    p = doc.add_paragraph()
    set_line_spacing(p, 1.5)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_bold_body(bold_text, normal_text=""):
    p = doc.add_paragraph()
    set_line_spacing(p, 1.5)
    run = p.add_run(bold_text)
    set_font(run, size=12, bold=True)
    if normal_text:
        run2 = p.add_run(normal_text)
        set_font(run2, size=12)
    return p

# ── Read markdown ──
with open('人工智能概论期末项目说明.md', 'r', encoding='utf-8') as f:
    md = f.read()

lines = md.split('\n')

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip HTML comments, separators, empty lines at top
    if not stripped:
        i += 1
        continue

    # ── Title (first heading) ──
    if stripped.startswith('# ') and i == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_line_spacing(p, 1.5)
        run = p.add_run(stripped[2:])
        set_font(run, size=18, bold=True)
        i += 1
        continue

    # ── Headings ──
    if stripped.startswith('## '):
        text = stripped[3:]
        # Determine level
        if text.startswith('选题'):
            add_heading_line(text, level=2)
        elif text in ('整体要求', '排版格式', '封面内容（第一页）'):
            add_heading_line(text, level=2)
        elif text.startswith('一、') or text.startswith('二、') or text.startswith('三、') or text.startswith('四、') or text.startswith('五、') or text.startswith('六、') or text.startswith('七、'):
            add_heading_line(text, level=1)
        else:
            add_heading_line(text, level=2)
        i += 1
        continue

    if stripped.startswith('### '):
        add_heading_line(stripped[4:], level=2)
        i += 1
        continue

    # ── Table ──
    if '|' in stripped and stripped.startswith('|'):
        # Collect table rows
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            rows.append(lines[i].strip())
            i += 1
        if len(rows) >= 2:
            # Skip separator row (|---|)
            data_rows = [r for r in rows if not re.match(r'^\|[\s\-:]+\|', r)]
            if data_rows:
                table = doc.add_table(rows=len(data_rows), cols=len(data_rows[0].split('|')) - 2)
                table.style = 'Table Grid'
                for row_idx, row_text in enumerate(data_rows):
                    cells = [c.strip() for c in row_text.split('|')][1:-1]
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            run = p.add_run(cell_text)
                            set_font(run, size=11)
                            # Bold first row
                            if row_idx == 0:
                                run.font.bold = True
            # add spacing after table
            doc.add_paragraph()
        continue

    # ── Rules (---) ──
    if stripped == '---':
        i += 1
        continue

    # ── Blockquote ──
    if stripped.startswith('> '):
        text = stripped[2:]
        p = doc.add_paragraph()
        set_line_spacing(p, 1.5)
        run = p.add_run(text)
        set_font(run, size=12, color=(100, 100, 100))
        i += 1
        continue

    if stripped.startswith('>'):
        text = stripped[1:]
        p = doc.add_paragraph()
        set_line_spacing(p, 1.5)
        run = p.add_run(text)
        set_font(run, size=12, color=(100, 100, 100))
        i += 1
        continue

    # ── Code block ──
    if stripped.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        if code_lines:
            p = doc.add_paragraph()
            set_line_spacing(p, 1.0)
            run = p.add_run(''.join(code_lines))
            set_font(run, name='Courier New', size=10)
        i += 1
        continue

    # ── Bullet / list ──
    if stripped.startswith('- '):
        text = stripped[2:]
        # Check for bold pattern: **text**
        p = doc.add_paragraph()
        set_line_spacing(p, 1.5)
        # Handle inline bold
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_font(run, size=12, bold=True)
            else:
                run = p.add_run(part)
                set_font(run, size=12)
        i += 1
        continue

    if re.match(r'^\d+\.\s', stripped):
        p = doc.add_paragraph()
        set_line_spacing(p, 1.5)
        run = p.add_run(stripped)
        set_font(run, size=12)
        i += 1
        continue

    # ── Regular body text ──
    # Handle inline bold
    if '**' in stripped:
        p = doc.add_paragraph()
        set_line_spacing(p, 1.5)
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_font(run, size=12, bold=True)
            else:
                run = p.add_run(part)
                set_font(run, size=12)
    else:
        add_body(stripped)

    i += 1

doc.save('人工智能概论期末项目说明.docx')
print('Done: 人工智能概论期末项目说明.docx')
